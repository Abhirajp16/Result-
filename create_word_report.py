from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_border(doc):
    """Add border to all pages"""
    section = doc.sections[0]
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '24')
        border_el.set(qn('w:space'), '24')
        border_el.set(qn('w:color'), '000000')
        pgBorders.append(border_el)
    
    sectPr.append(pgBorders)

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)

# Add page border
add_page_border(doc)

# ============================================================================
# TITLE PAGE
# ============================================================================

# VTU Header
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('VISVESVARAYA TECHNOLOGICAL UNIVERSITY')
run.font.size = Pt(16)
run.font.bold = True
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Jnana Sangama, Belagavi – 590 014')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROJECT PHASE I (BCS685)')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('on')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AUTOMATED VTU RESULT FETCHING AND ANALYSIS SYSTEM')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Submitted in partial fulfillment for the award of the degree of Bachelor of Engineering')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('in')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('COMPUTER SCIENCE AND ENGINEERING')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Submitted By')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[STUDENT 1 NAME] ([USN])')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[STUDENT 2 NAME] ([USN])')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[STUDENT 3 NAME] ([USN])')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[STUDENT 4 NAME] ([USN])')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Under the guidance of')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[GUIDE NAME], M.Tech., PhD.')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Assistant Professor, Department of CSE, [COLLEGE NAME]')
run.font.size = Pt(11)
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[INSERT COLLEGE LOGO HERE]')
run.font.size = Pt(10)
run.font.italic = True
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[YOUR COLLEGE NAME]')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BANGALORE-560 048')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2025-26')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

# Page break
doc.add_page_break()

# ============================================================================
# CERTIFICATE PAGE
# ============================================================================

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[YOUR COLLEGE NAME]')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[ISO Certified 9002:2008, Affiliated to VTU, Belagavi, Approved by AICTE, New Delhi]')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[College Address]')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CERTIFICATE')
run.font.size = Pt(16)
run.font.bold = True
run.font.name = 'Times New Roman'
run.underline = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Certified that the project work entitled ')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run = p.add_run('"AUTOMATED VTU RESULT FETCHING AND ANALYSIS SYSTEM"')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = 'Times New Roman'
run = p.add_run(' carried out by ')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run = p.add_run('[STUDENT 1 NAME] ([USN]), [STUDENT 2 NAME] ([USN]), [STUDENT 3 NAME] ([USN]), [STUDENT 4 NAME] ([USN])')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = 'Times New Roman'
run = p.add_run(', bonafide students of ')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run = p.add_run('[YOUR COLLEGE NAME]')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = 'Times New Roman'
run = p.add_run(' in partial fulfilment for the award of Bachelor of Engineering in Computer Science and Engineering of the Visvesvaraya Technological University, Belgaum during the year, 2025-2026. It is certified that all suggestions indicated for Internal Assessment have been incorporated in the Report deposited in the Department library. The project report has been approved as it satisfies the academic requirements in respect of Project work prescribed for the said degree.')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Signature table
table = doc.add_table(rows=2, cols=3)
table.style = 'Table Grid'

# Row 1
table.rows[0].cells[0].text = 'Signature of the Guide'
table.rows[0].cells[1].text = 'Signature of the HOD'
table.rows[0].cells[2].text = 'Signature of the Principal'

# Row 2
table.rows[1].cells[0].text = '[Guide Name]'
table.rows[1].cells[1].text = '[HOD Name]'
table.rows[1].cells[2].text = '[Principal Name]'

# Format table
for row in table.rows:
    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'

doc.add_page_break()

# ============================================================================
# ABSTRACT
# ============================================================================

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ABSTRACT')
run.font.size = Pt(16)
run.font.bold = True
run.font.name = 'Times New Roman'
run.underline = True

doc.add_paragraph()

abstract_text = """Academic result management in universities involves repetitive manual processes where students must navigate CAPTCHA verifications and manually track performance across semesters. This project presents an automated system for fetching and analyzing VTU examination results with intelligent performance tracking features.

The system uses a deep learning model combining Convolutional Neural Networks (CNN) and Bidirectional Gated Recurrent Units (BiGRU) with Connectionist Temporal Classification (CTC) loss to automatically solve CAPTCHA challenges. Selenium WebDriver automates browser interaction and data extraction. Results are stored in a database enabling comprehensive analysis across all eight semesters.

Key features include role-based access for students, teachers, and HODs, automatic SGPA and CGPA calculation, revaluation tracking with before-and-after comparisons, top performer identification, class average computation, and visual dashboards. The system generates PDF reports and supports concurrent multi-user access through a Flask-SocketIO web interface with real-time updates.

This solution reduces manual effort in result checking and provides valuable insights into student performance trends, helping students track progress, teachers identify struggling students, and administrators make data-driven decisions for academic planning."""

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run(abstract_text)
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Keywords: ')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = 'Times New Roman'
run = p.add_run('VTU Result Automation, CAPTCHA Recognition, Deep Learning, Academic Performance Analysis, Web Scraping, SGPA-CGPA Calculation')
run.font.size = Pt(12)
run.font.italic = True
run.font.name = 'Times New Roman'

doc.add_page_break()

# ============================================================================
# CHAPTER 1: INTRODUCTION
# ============================================================================

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CHAPTER 1')
run.font.size = Pt(16)
run.font.bold = True
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('INTRODUCTION')
run.font.size = Pt(16)
run.font.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph()

intro_text = """The digitization of educational systems has transformed how academic information is accessed. Universities now publish examination results online, but challenges remain in result management and performance analysis. Students face repetitive manual processes, lack of historical tracking, and limited analytical insights.

This project develops an intelligent automated system for VTU result fetching and analysis. By combining deep learning for CAPTCHA recognition, web automation for data extraction, and comprehensive analytics, the system transforms how students, teachers, and administrators interact with academic results, enabling data-driven decisions in education."""

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run(intro_text)
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

# 1.1 BACKGROUND
p = doc.add_paragraph()
run = p.add_run('1.1 BACKGROUND')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

background_text = """VTU publishes examination results through an online portal requiring CAPTCHA verification for each query. This becomes tedious when checking results multiple times or tracking performance across semesters. Traditional systems lack analytics, forcing students to manually calculate SGPA/CGPA while teachers spend hours compiling class data.

Recent advances in deep learning enable automated CAPTCHA recognition. CNN-RNN models can read text from images, while Selenium automates web interaction, allowing an intelligent system that fetches results and provides meaningful analytics."""

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run(background_text)
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

# 1.2 AIM
p = doc.add_paragraph()
run = p.add_run('1.2 AIM OF THE PROJECT')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

aim_text = """To develop an automated system for fetching VTU examination results and providing comprehensive academic performance analysis through deep learning-based CAPTCHA solving and role-based dashboards for actionable insights."""

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run(aim_text)
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

# 1.3 PROBLEM STATEMENT
p = doc.add_paragraph()
run = p.add_run('1.3 PROBLEM STATEMENT')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

problem_text = """Current VTU result management faces repetitive manual CAPTCHA solving, lack of historical tracking, error-prone manual calculations, difficulty tracking revaluation changes, limited analytics, no multi-user access, and automation barriers. These create inefficiency and prevent data-driven academic decisions."""

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run(problem_text)
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

# 1.4 OBJECTIVES
p = doc.add_paragraph()
run = p.add_run('1.4 OBJECTIVES')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

objectives = [
    "Automate CAPTCHA recognition using CNN-BiGRU-CTC model",
    "Automate result fetching through Selenium web scraping",
    "Build centralized database for all semester results",
    "Implement role-based access for students, teachers, and HODs",
    "Automatically calculate SGPA and CGPA",
    "Track revaluation with before-and-after comparison",
    "Identify top 5 students before and after revaluation",
    "Compute class average SGPA",
    "Provide semester 1-8 analysis",
    "Create visual dashboards with charts",
    "Generate downloadable PDF reports",
    "Enable real-time updates using Flask-SocketIO"
]

for i, obj in enumerate(objectives, 1):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(obj)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

doc.add_page_break()

# Save document
doc.save('VTU_Result_Analysis_Report.docx')
print("✅ Word document created successfully!")
print("📄 File saved as: VTU_Result_Analysis_Report.docx")
print("\n⚠️ Note: You need to install python-docx library first:")
print("   Run: pip install python-docx")
